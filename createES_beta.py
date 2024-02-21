import Status
import Company
from dotenv import load_dotenv

load_dotenv()

from langchain_google_genai import ChatGoogleGenerativeAI

llm = ChatGoogleGenerativeAI(model='gemini-pro', temperature=0)

status = Status.Status()
company = Company.Company()

import openpyxl

wb = openpyxl.load_workbook(r"C:\Users\bothm\dev_Python\createES_RPA\create_entry_sheet_beta.xlsm")
company_sheet = wb['createES']
status_sheet = wb['Status']

company.setName(company_sheet['C4'].value)
company.setJob(company_sheet['C8'].value)
company.setQuestion(company_sheet['C15'].value)
status.setCharactorLimit(company_sheet['C12'].value)
status.setStrength(status_sheet['C4'].value)
status.setExperience(status_sheet['C8'].value, status_sheet['C12'].value)

strengthPrompt = status.getStrength(0)
for strengthNum in range(1, status.getStrengthNum()):
    strengthPrompt += "、" + status.getStrength(strengthNum)

experienceKeys = list(status.experience.keys())
experiencePrompt = ""
for experienceNum in range(0, len(experienceKeys)):
    experiencePrompt += experienceKeys[experienceNum] + "で身に着けた" + status.experience.get(experienceKeys[experienceNum]) + "\n"

firstPrompt = "あなたは素晴らしいAIです。次の質問に日本語で答えてください。"
mainPrompt = (
    company.getName()
    + "の"
    + company.getJob()
    + "職の公式HPを参照し、"
    + "「" + str(company.getQuestion()) + "」という質問に対しての回答を"
    + str(status.getCharactorLimit())
    + "以内で回答してください。会社名を言うときは、「貴社」という言葉を使用してください。また、下記が私の強みと経験です。これらの強みと経験が、企業のどのような部分とマッチするかを明確にしてください。なお、回答の文字数は"
    + str(status.getCharactorLimit())
    + "字の9割以上にする必要があります。企業の理念やミッションを引用するときは、HPに記載されている内容をそのまま使用してください。\n"
    + "強み： " + strengthPrompt
    + "\n経験： " + experiencePrompt
)

result = llm.invoke(str(mainPrompt))

# Excelに保存
# company_sheet['J4'] = result.content.replace("\n", "")
# company_sheet['J4'].alignment = openpyxl.styles.Alignment(wrapText=True)
# wb.save(r"C:\Users\bothm\dev_Python\createES_RPA\{}.xlsx".format(company_sheet['C4'].value))

# Wordに保存(ここから)
from docx import Document
from docx.shared import Pt
import os

target_document = r"C:\Users\bothm\dev_Python\createES_RPA\{}.docx".format(company_sheet['C4'].value)

# ファイルが既に存在するか
if os.path.exists(target_document):
    # ドキュメントオブジェクト作成
    document = Document(target_document)

    # ファイルに追記
    document.add_paragraph("")
    document.add_paragraph(str(company.getQuestion()))
    document.add_paragraph(str(result.content))
    document.paragraphs[len(document.paragraphs) - 2].runs[0].font.size = Pt(12)
    document.paragraphs[len(document.paragraphs) - 2].runs[0].bold = True

    # 上書き保存
    document.save(target_document) 
else:
    # ドキュメントオブジェクト作成
    document = Document()

    # ファイル作成
    document.add_paragraph(str(company.getQuestion()))
    document.add_paragraph(str(result.content))
    document.paragraphs[0].runs[0].font.size = Pt(12)
    document.paragraphs[0].runs[0].bold = True

    # 新規保存
    document.save(target_document)